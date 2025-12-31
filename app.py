"""Flask web app for Short Story Pipeline."""

# Load environment variables from .env file before other imports
from dotenv import load_dotenv  # type: ignore[import-untyped]

load_dotenv()  # noqa: E402

import os  # noqa: E402
import uuid  # noqa: E402
import re  # noqa: E402
import logging  # noqa: E402
from io import BytesIO  # noqa: E402
from datetime import datetime  # noqa: E402
from urllib.parse import quote  # noqa: E402
from flask import Flask, render_template, request, jsonify, send_file, Response  # noqa: E402
from flask_cors import CORS  # type: ignore[import-untyped]  # noqa: E402
from flask_limiter import Limiter  # type: ignore[import-untyped]  # noqa: E402
from flask_limiter.util import get_remote_address  # type: ignore[import-untyped]  # noqa: E402
from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]  # noqa: E402
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore[import-untyped]  # noqa: E402
from reportlab.lib.units import inch  # type: ignore[import-untyped]  # noqa: E402
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore[import-untyped]  # noqa: E402
from reportlab.lib.enums import TA_LEFT  # type: ignore[import-untyped]  # noqa: E402
from reportlab.lib.colors import HexColor  # type: ignore[import-untyped]  # noqa: E402
from src.shortstory.pipeline import ShortStoryPipeline  # noqa: E402
from src.shortstory.genres import get_available_genres, get_genre_config  # noqa: E402
from src.shortstory.templates import (  # noqa: E402
    get_templates_for_genre, get_all_templates, get_available_template_genres
)
from src.shortstory.utils import (  # noqa: E402
    check_distinctiveness, MAX_WORD_COUNT,
    create_story_repository,
)
from src.shortstory.memorability_scorer import get_memorability_scorer  # noqa: E402
from src.shortstory.utils.errors import (  # noqa: E402
    register_error_handlers, ValidationError, NotFoundError, ServiceUnavailableError,
    MissingDependencyError
)  # noqa: E402

# Configure logging
logging.basicConfig(
    level=logging.INFO if os.getenv('FLASK_ENV') != 'development' else logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

# Configure rate limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri=os.getenv('REDIS_URL', 'memory://'),  # Use Redis in production if available
    headers_enabled=True
)

pipeline = ShortStoryPipeline()

# Initialize story repository (unified storage interface)
story_repository = create_story_repository()


def init_stories():
    """
    Initialize story storage on application startup.
    
    Logs the current story count from the repository.
    """
    count = story_repository.count()
    logger.info(f"Story repository initialized with {count} stories")


def check_llm_setup():
    """
    Check LLM (Large Language Model) setup and provide helpful feedback.
    
    Verifies that the Google Gemini API is properly configured by checking:
    1. Whether the GOOGLE_API_KEY environment variable is set
    2. Whether the API client can be initialized
    3. Whether the configured model is available and accessible
    
    Prints status messages to the console indicating the setup status.
    This function is typically called during application startup to verify
    that AI-powered story generation will be available.
    
    Note:
        If the API is not available, the application will fall back to
        template-based story generation instead of AI generation.
    """
    try:
        from src.shortstory.utils import get_default_client
        import os
        
        has_api_key = bool(os.getenv("GOOGLE_API_KEY"))
        
        if has_api_key:
            try:
                client = get_default_client()
                is_available = client.check_availability()
                
                if is_available:
                    print(f"✅ Google Gemini API ready: Using model '{client.model_name}'")
                    print("   Story generation will use AI-powered prose generation")
                else:
                    print(f"⚠️  Google Gemini API configured but model '{client.model_name}' may not be available")
                    print("   The app will try to use AI generation, but may fall back to templates")
            except ValueError as e:
                print(f"⚠️  {e}")
                print("   The app will work with template-based generation")
                print("   See SETUP_GOOGLE.md for setup instructions")
        else:
            print("⚠️  GOOGLE_API_KEY not set")
            print("   The app will work, but will use template-based generation")
            print("   To enable AI generation:")
            print("   1. Get API key from https://makersuite.google.com/app/apikey")
            print("   2. Set GOOGLE_API_KEY environment variable")
            print("   See SETUP_GOOGLE.md for details")
    except Exception as e:
        print(f"⚠️  Could not check LLM setup: {e}")
        print("   The app will work with template-based generation")
        print("   See SETUP_GOOGLE.md for setup instructions")


# Register error handlers
register_error_handlers(app, debug=os.getenv('FLASK_ENV') == 'development')

# Initialize on startup
init_stories()
check_llm_setup()


def word_count_response(word_count, max_words=MAX_WORD_COUNT):
    """
    Build standardized word count response.
    
    Args:
        word_count: The current word count of the story
        max_words: The maximum allowed word count (default: MAX_WORD_COUNT)
    
    Returns:
        dict: A dictionary containing word count information with keys:
            - word_count: Current word count
            - max_words: Maximum allowed words
            - remaining_words: Words remaining before limit
    """
    return {
        "word_count": word_count,
        "max_words": max_words,
        "remaining_words": max_words - word_count
    }


def get_story_or_404(story_id):
    """
    Get story from repository, or return None.
    
    Args:
        story_id: Unique identifier for the story
    
    Returns:
        dict: Story data dictionary if found, None otherwise
    """
    return story_repository.load(story_id)


def sanitize_filename(title, story_id, max_length=50):
    """
    Sanitize a title for use as a safe filename.
    
    Removes potentially dangerous characters that could lead to XSS vulnerabilities
    or filesystem issues. Limits the filename length and provides a fallback if
    sanitization removes all characters.
    
    Args:
        title: The original title to sanitize
        story_id: Story identifier to use in fallback filename
        max_length: Maximum length for the sanitized filename (default: 50)
    
    Returns:
        str: A safe filename string with dangerous characters removed, spaces
             replaced with underscores, and length limited to max_length.
             Falls back to "story_{story_id}" if sanitization removes all characters.
    
    Example:
        >>> sanitize_filename("My Story<script>", "abc123")
        'My_Story'
        >>> sanitize_filename("<>\"'", "abc123")
        'story_abc123'
    """
    # Remove dangerous characters that could lead to XSS or filesystem issues
    # Characters removed: < > " ' ; \ / and control characters (\r\n\t)
    safe_filename = re.sub(r'[<>"\';\\/\r\n\t]', '', title.replace(' ', '_'))
    # Remove directory traversal sequences
    safe_filename = safe_filename.replace('..', '')
    safe_filename = safe_filename[:max_length]  # Limit length
    
    # Fallback if sanitization removes everything
    if not safe_filename:
        safe_filename = f"story_{story_id}"
    
    return safe_filename


@app.route('/')
def index():
    return render_template('index.html', genres=get_available_genres())


@app.route('/api/health')
def health():
    return jsonify({"status": "ok"})


@app.route('/api/genres', methods=['GET'])
def get_genres():
    return jsonify({"genres": get_available_genres()})


@app.route('/api/generate', methods=['POST'])
@limiter.limit("10 per minute")  # Limit story generation to 10 per minute per IP
def generate_story():
    try:
        data = request.get_json() or {}
        idea = data.get('idea', '').strip()
        character = data.get('character', {})
        theme = data.get('theme', '').strip()
        genre = data.get('genre', 'General Fiction')
        
        if not idea:
            raise ValidationError(
                "Story idea is required. Please provide a creative premise for your story.",
                details={"field": "idea"}
            )
        
        if isinstance(character, str):
            character = {"description": character}
        
        genre_config = get_genre_config(genre)
        if genre_config is None:
            # Fallback to General Fiction if genre not found
            genre_config = get_genre_config('General Fiction')
        # Ensure genre_config is not None (should always be a dict)
        if genre_config is None:
            raise ServiceUnavailableError("genre_config", "Genre configuration service unavailable")
        
        pipeline.genre = genre
        pipeline.genre_config = genre_config
        
        premise = pipeline.capture_premise(idea, character, theme, validate=True)
        outline = pipeline.generate_outline(genre=genre)
        scaffold = pipeline.scaffold(genre=genre)
        
        # Run full pipeline: draft and revise
        draft = pipeline.draft()
        revised_draft = pipeline.revise()
        
        framework = genre_config.get('framework', 'narrative_arc')
        outline_structure = genre_config.get('outline', ['setup', 'complication', 'resolution'])
        constraints = genre_config.get('constraints', {})
        tone = scaffold.get('tone', constraints.get('tone', 'balanced')) if isinstance(scaffold, dict) else constraints.get('tone', 'balanced')
        pace = scaffold.get('pace', constraints.get('pace', 'moderate')) if isinstance(scaffold, dict) else constraints.get('pace', 'moderate')
        pov = scaffold.get('pov', constraints.get('pov_preference', 'flexible')) if isinstance(scaffold, dict) else constraints.get('pov_preference', 'flexible')
        
        idea_dist = check_distinctiveness(idea)
        char_dist = check_distinctiveness(None, character=character)
        
        char_desc = character.get('description', character) if isinstance(character, dict) else character
        cliche_msg = f"⚠️ Clichés: {', '.join(idea_dist.get('found_cliches', []))}" if idea_dist.get('has_cliches') else "✓ No clichés"
        generic_msg = f"⚠️ Generic: {', '.join(char_dist.get('generic_elements', []))}" if char_dist.get('has_generic_archetype') else "✓ Distinctive"
        
        # Get the revised story text
        revised_story_text = revised_draft.get('text', '')
        
        # Build the full story output with metadata
        story_text = f"""# {idea}

## Genre: {genre} | Framework: {framework} | Structure: {', '.join(outline_structure)}
**Tone:** {tone} | **Pace:** {pace} | **POV:** {pov}

## Character
{char_desc if char_desc else 'Not specified'}

## Theme
{theme if theme else 'Not specified'}

## Distinctiveness
Idea: {idea_dist['distinctiveness_score']:.2f}/1.0 | Character: {char_dist['distinctiveness_score']:.2f}/1.0
{cliche_msg}
{generic_msg}

## Story

{revised_story_text}

**Constraints:** {', '.join([f"{k}:{v}" for k, v in constraints.items()]) if constraints else 'None specified'}
"""
        
        # Use word count from revised draft
        word_count = revised_draft.get('word_count', 0)
        story_id = f"story_{uuid.uuid4().hex[:8]}"
        # Initialize revision history
        revision_history = [{
            "version": 1,
            "text": draft.get('text', ''),
            "word_count": draft.get('word_count', 0),
            "type": "draft",
            "timestamp": datetime.now().isoformat()
        }, {
            "version": 2,
            "text": revised_draft.get('text', ''),
            "word_count": revised_draft.get('word_count', 0),
            "type": "revised",
            "timestamp": datetime.now().isoformat()
        }]
        
        story_data = {
            "id": story_id, "premise": premise, "outline": outline, "genre": genre,
            "genre_config": genre_config, "text": story_text,
            "word_count": word_count, "max_words": MAX_WORD_COUNT,
            "draft": draft, "revised_draft": revised_draft,
            "revision_history": revision_history,
            "current_revision": 2
        }
        
        # Save to repository
        story_repository.save(story_data)
        
        logger.info(f"Successfully generated story {story_id}")
        return jsonify({
            "success": True, "story_id": story_id, "story": story_text,
            **word_count_response(word_count), "premise": premise,
            "outline": outline, "genre": genre, "genre_config": genre_config
        })
    except ValidationError:
        # Re-raise validation errors to be handled by error handler
        raise
    except ValueError as e:
        error_msg = str(e)
        if "premise" in error_msg.lower() or "validation" in error_msg.lower():
            raise ValidationError(
                f"Validation error: {error_msg}. Please refine your story idea or character description.",
                details={"field": "premise", "original_error": error_msg}
            )
        raise ValidationError(
            f"Invalid input: {error_msg}. Please check your story parameters.",
            details={"original_error": error_msg}
        )
    except Exception as e:
        error_msg = str(e)
        # Check for common API errors
        if "api" in error_msg.lower() or "key" in error_msg.lower():
            logger.warning(f"API error during story generation: {error_msg}")
            raise ServiceUnavailableError(
                "ai_generation",
                "AI generation service error. The app will use template-based generation. Check your API configuration if you expected AI generation."
            )
        # Re-raise to be handled by generic error handler
        logger.error(f"Unexpected error during story generation: {error_msg}", exc_info=True)
        raise


@app.route('/api/story/<story_id>', methods=['GET'])
@limiter.limit("100 per hour")
def get_story(story_id):
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    word_count = pipeline.word_validator.count_words(story["text"])
    return jsonify({
        "story_id": story_id, "story": story["text"],
        **word_count_response(word_count, story["max_words"]),
        "premise": story.get("premise")
    })


@app.route('/api/story/<story_id>', methods=['PUT'])
@limiter.limit("30 per hour")
def update_story(story_id):
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    
    data = request.get_json()
    if not data or 'text' not in data:
        raise ValidationError(
            "Story text is required in the request body.",
            details={"field": "text"}
        )
    
    word_count, is_valid = pipeline.word_validator.validate(data['text'], raise_error=False)
    if not is_valid:
        raise ValidationError(
            f"Story exceeds word limit ({word_count:,} > {MAX_WORD_COUNT:,} words). Please reduce the length.",
            details={
                "word_count": word_count,
                "max_words": MAX_WORD_COUNT,
                **word_count_response(word_count)
            }
        )
    
    story["text"] = data['text']
    story["word_count"] = word_count
    
    # Update in repository
    story_repository.update(story_id, {"text": data['text'], "word_count": word_count})
    
    return jsonify({"success": True, "story_id": story_id, **word_count_response(word_count)})


@app.route('/api/validate', methods=['POST'])
@limiter.limit("100 per hour")
def validate_text():
    data = request.get_json()
    if not data or 'text' not in data:
        raise ValidationError(
            "Text content is required for validation.",
            details={"field": "text"}
        )
    
    word_count, is_valid = pipeline.word_validator.validate(data['text'], raise_error=False)
    response = {
        **word_count_response(word_count),
        "is_valid": is_valid,
        "distinctiveness": check_distinctiveness(data['text'])
    }
    
    # Optionally include memorability score if requested
    if data.get('include_memorability', False):
        scorer = get_memorability_scorer()
        character = data.get('character')
        outline = data.get('outline')
        premise = data.get('premise')
        response['memorability'] = scorer.score_story(
            text=data['text'],
            character=character,
            outline=outline,
            premise=premise
        )
    
    return jsonify(response)


@app.route('/api/memorability/score', methods=['POST'])
@limiter.limit("50 per hour")
def score_memorability():
    """
    Score story memorability across multiple dimensions.
    
    Request body:
    {
        "text": str (required),
        "character": dict (optional),
        "outline": dict (optional),
        "premise": dict (optional)
    }
    
    Returns:
    {
        "overall_score": float,
        "dimensions": {...},
        "prioritized_suggestions": [...],
        "summary": str,
        "detailed_analysis": {...}
    }
    """
    data = request.get_json()
    if not data or 'text' not in data:
        raise ValidationError(
            "Text content is required for memorability scoring.",
            details={"field": "text"}
        )
    
    scorer = get_memorability_scorer()
    result = scorer.score_story(
        text=data['text'],
        character=data.get('character'),
        outline=data.get('outline'),
        premise=data.get('premise')
    )
    
    return jsonify(result)


@app.route('/api/story/<story_id>/save', methods=['POST'])
@limiter.limit("30 per hour")
def save_story_endpoint(story_id):
    """Explicitly save a story to disk. Does NOT create new revisions."""
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    
    # Optionally update text if provided
    data = request.get_json() or {}
    if 'text' in data:
        word_count, is_valid = pipeline.word_validator.validate(data['text'], raise_error=False)
        if not is_valid:
            raise ValidationError(
                f"Story exceeds word limit ({word_count:,} > {MAX_WORD_COUNT:,} words). Please reduce the length.",
                details={
                    "word_count": word_count,
                    "max_words": MAX_WORD_COUNT,
                    **word_count_response(word_count)
                }
            )
        story["text"] = data['text']
        story["word_count"] = word_count
    
    # IMPORTANT: Save does NOT create new revisions - it just updates the current story
    # Only the /revise endpoint creates new revisions
    
    # Save to repository
    success = story_repository.save(story)
    
    if success:
        logger.info(f"Story {story_id} saved successfully (no new revision created)")
        return jsonify({
            "success": True,
            "message": "Story saved successfully",
            "story_id": story_id,
            "word_count": story.get("word_count", 0)
        })
    else:
        logger.error(f"Failed to save story {story_id}")
        raise ServiceUnavailableError("storage", "Failed to save story. Please try again.")


@app.route('/api/stories', methods=['GET'])
@limiter.limit("100 per hour")
def list_all_stories():
    """
    List all stories with metadata.
    
    Supports pagination via query parameters:
    - page: Page number (default: 1)
    - per_page: Items per page (default: 50, max: 100)
    
    Returns:
        JSON response with paginated stories, total count, and pagination metadata
    """
    try:
        # Get pagination parameters
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 50, type=int)
        genre = request.args.get('genre', None, type=str)
        
        # Validate and limit per_page
        per_page = min(max(1, per_page), 100)  # Between 1 and 100
        page = max(1, page)  # At least page 1
        
        # Use repository for listing (handles pagination internally)
        result = story_repository.list(page=page, per_page=per_page, genre=genre)
        return jsonify({
            "success": True,
            **result
        })
    except Exception as e:
        logger.error(f"Failed to load stories: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("storage", "Failed to load stories. Please try again.")


@app.route('/api/templates', methods=['GET'])
def get_templates():
    """Get story templates, optionally filtered by genre."""
    genre = request.args.get('genre')
    if genre:
        templates = get_templates_for_genre(genre)
        return jsonify({
            "success": True,
            "genre": genre,
            "templates": templates
        })
    else:
        all_templates = get_all_templates()
        return jsonify({
            "success": True,
            "templates": all_templates,
            "genres": get_available_template_genres()
        })


@app.route('/api/story/<story_id>/revise', methods=['POST'])
@limiter.limit("20 per hour")
def revise_story(story_id):
    """Run an additional revision pass on a story."""
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    
    # Get current story text
    current_text = story.get("text", "")
    if not current_text:
        raise ValidationError(
            "Story has no content to revise.",
            details={"story_id": story_id}
        )
    
    try:
        # Create a temporary draft object for revision
        temp_draft = {
            "text": current_text,
            "word_count": story.get("word_count", 0)
        }
        
        # Run revision
        revised_draft = pipeline.revise(draft=temp_draft, use_llm=True)
        revised_text = revised_draft.get('text', '')
        revised_word_count = revised_draft.get('word_count', 0)
        
        # Update story with revised text
        story["text"] = revised_text
        story["word_count"] = revised_word_count
        
        # Add to revision history
        if "revision_history" not in story:
            story["revision_history"] = []
        if "current_revision" not in story:
            story["current_revision"] = 0
        
        new_version = story["current_revision"] + 1
        story["revision_history"].append({
            "version": new_version,
            "text": revised_text,
            "word_count": revised_word_count,
            "type": "revised",
            "timestamp": datetime.now().isoformat()
        })
        story["current_revision"] = new_version
        story["revised_draft"] = revised_draft
        
        # Save to repository
        story_repository.save(story)
        
        logger.info(f"Story {story_id} revised to version {new_version}")
        return jsonify({
            "success": True,
            "story_id": story_id,
            "story": revised_text,
            "revision_number": new_version,
            **word_count_response(revised_word_count)
        })
    except Exception as e:
        logger.error(f"Revision failed for story {story_id}: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("revision", f"Revision failed: {str(e)}")


@app.route('/api/story/<story_id>/revisions', methods=['GET'])
@limiter.limit("100 per hour")
def get_revision_history(story_id):
    """Get revision history for a story."""
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    
    revision_history = story.get("revision_history", [])
    current_revision = story.get("current_revision", len(revision_history))
    
    return jsonify({
        "success": True,
        "story_id": story_id,
        "revision_history": revision_history,
        "current_revision": current_revision,
        "total_revisions": len(revision_history)
    })


@app.route('/api/story/<story_id>/compare', methods=['GET'])
@limiter.limit("100 per hour")
def compare_story_versions(story_id):
    """Compare two versions of a story."""
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    
    revision_history = story.get("revision_history", [])
    if len(revision_history) < 2:
        raise ValidationError(
            "Not enough revisions to compare. Need at least 2 versions.",
            details={"story_id": story_id, "revision_count": len(revision_history)}
        )
    
    version1 = request.args.get('version1', type=int)
    version2 = request.args.get('version2', type=int)
    
    if version1 is None or version2 is None:
        # Default to first and last
        version1 = 1
        version2 = len(revision_history)
    
    # Find versions
    v1_data = next((r for r in revision_history if r.get("version") == version1), None)
    v2_data = next((r for r in revision_history if r.get("version") == version2), None)
    
    if not v1_data or not v2_data:
        available_versions = [r.get('version') for r in revision_history]
        raise ValidationError(
            f"Version(s) not found. Available versions: {available_versions}",
            details={
                "story_id": story_id,
                "requested_versions": {"version1": version1, "version2": version2},
                "available_versions": available_versions
            }
        )
    
    # Simple diff calculation (word-level for now)
    text1 = v1_data.get("text", "")
    text2 = v2_data.get("text", "")
    
    words1 = text1.split()
    words2 = text2.split()
    
    # Calculate basic statistics
    word_count_diff = v2_data.get("word_count", 0) - v1_data.get("word_count", 0)
    
    return jsonify({
        "success": True,
        "story_id": story_id,
        "version1": {
            "version": version1,
            "text": text1,
            "word_count": v1_data.get("word_count", 0),
            "timestamp": v1_data.get("timestamp"),
            "type": v1_data.get("type", "unknown")
        },
        "version2": {
            "version": version2,
            "text": text2,
            "word_count": v2_data.get("word_count", 0),
            "timestamp": v2_data.get("timestamp"),
            "type": v2_data.get("type", "unknown")
        },
        "comparison": {
            "word_count_diff": word_count_diff,
            "text_length_diff": len(text2) - len(text1),
            "words_added": max(0, len(words2) - len(words1)),
            "words_removed": max(0, len(words1) - len(words2))
        }
    })


@app.route('/api/story/<story_id>/export/<format_type>', methods=['GET'])
@limiter.limit("50 per hour")
def export_story(story_id, format_type):
    """Export a story in various formats (pdf, markdown, txt)."""
    story = get_story_or_404(story_id)
    if story is None:
        raise NotFoundError("Story", story_id)
    
    story_text = story.get("text", "")
    if not story_text:
        raise ValidationError(
            "Story has no content to export",
            details={"story_id": story_id}
        )
    
    # Extract title from story text (first line after #)
    title_match = re.search(r'^#\s+(.+)$', story_text, re.MULTILINE)
    raw_title = title_match.group(1) if title_match else f"Story {story_id}"
    
    # Sanitize title to prevent XSS vulnerabilities
    # Remove or escape potentially malicious characters
    title = re.sub(r'[<>"\';\\/]', '', raw_title)  # Remove dangerous characters
    title = title.strip()[:100]  # Limit length and trim whitespace
    if not title:  # Fallback if sanitization removes everything
        title = f"Story_{story_id}"
    
    try:
        if format_type == 'pdf':
            return export_pdf(story_text, title, story_id)
        elif format_type == 'markdown':
            return export_markdown(story_text, title, story_id)
        elif format_type == 'txt':
            return export_txt(story_text, title, story_id)
        elif format_type == 'docx':
            return export_docx(story_text, title, story_id)
        elif format_type == 'epub':
            return export_epub(story_text, title, story_id)
        else:
            raise ValidationError(
                f"Unsupported export format: {format_type}. Supported formats: pdf, markdown, txt, docx, epub",
                details={"format": format_type, "supported_formats": ["pdf", "markdown", "txt", "docx", "epub"]}
            )
    except (ValidationError, MissingDependencyError, ServiceUnavailableError):
        # Re-raise known error types to be handled by error handlers
        raise
    except Exception as e:
        logger.error(f"Export failed for story {story_id}, format {format_type}: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"Export failed: {str(e)}")


def export_pdf(story_text, title, story_id):
    """
    Export story as PDF format.
    
    Args:
        story_text: The story text to export (may contain markdown)
        title: The story title (will be sanitized for filename)
        story_id: Unique identifier for the story
    
    Returns:
        Flask send_file response with PDF attachment
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=HexColor('#667eea'),
        spaceAfter=30,
        alignment=TA_LEFT
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=12,
        alignment=TA_LEFT
    )
    
    story = []
    
    # Add title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Process story text - convert markdown headers to paragraphs
    lines = story_text.split('\n')
    for line in lines:
        if line.strip():
            # Remove markdown formatting for PDF
            clean_line = re.sub(r'^#+\s+', '', line)  # Remove headers
            clean_line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean_line)  # Bold
            clean_line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', clean_line)  # Italic
            story.append(Paragraph(clean_line, body_style))
        else:
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    buffer.seek(0)
    
    # Sanitize filename for safe download
    safe_filename = sanitize_filename(title, story_id, max_length=50)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f"{safe_filename}_{story_id}.pdf"
    )


def export_markdown(story_text, title, story_id):
    """
    Export story as Markdown format.
    
    Args:
        story_text: The story text to export
        title: The story title (will be sanitized for filename)
        story_id: Unique identifier for the story
    
    Returns:
        Flask Response with Markdown content and attachment headers
    """
    # Sanitize filename for safe download
    safe_filename = sanitize_filename(title, story_id, max_length=50)
    
    # Use RFC 5987 encoding for safer filename handling in HTTP headers
    # This prevents XSS attacks through filename injection
    encoded_filename = quote(safe_filename, safe='')
    filename_header = f'attachment; filename="{safe_filename}_{story_id}.md"; filename*=UTF-8\'\'{encoded_filename}_{story_id}.md'
    
    return Response(
        story_text,
        mimetype='text/markdown',
        headers={
            'Content-Disposition': filename_header
        }
    )


def export_txt(story_text, title, story_id):
    """
    Export story as plain text format.
    
    Removes markdown formatting and exports as plain text.
    
    Args:
        story_text: The story text to export (markdown will be removed)
        title: The story title (will be sanitized for filename)
        story_id: Unique identifier for the story
    
    Returns:
        Flask Response with plain text content and attachment headers
    """
    # Remove markdown formatting
    text = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    
    # Sanitize filename for safe download
    safe_filename = sanitize_filename(title, story_id, max_length=50)
    
    # Use RFC 5987 encoding for safer filename handling in HTTP headers
    # This prevents XSS attacks through filename injection
    encoded_filename = quote(safe_filename, safe='')
    filename_header = f'attachment; filename="{safe_filename}_{story_id}.txt"; filename*=UTF-8\'\'{encoded_filename}_{story_id}.txt'
    
    return Response(
        text,
        mimetype='text/plain',
        headers={
            'Content-Disposition': filename_header
        }
    )


def export_docx(story_text, title, story_id):
    """
    Export story as DOCX (Microsoft Word) format.
    
    Requires python-docx library. Raises MissingDependencyError if not installed.
    
    Args:
        story_text: The story text to export (may contain markdown)
        title: The story title (will be sanitized for filename)
        story_id: Unique identifier for the story
    
    Returns:
        Flask send_file response with DOCX attachment
    
    Raises:
        MissingDependencyError: If python-docx is not installed
        ServiceUnavailableError: If export fails for other reasons
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt  # type: ignore[import-untyped]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    except ImportError:
        raise MissingDependencyError("python-docx", "pip install python-docx")
    
    try:
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process story text
        lines = story_text.split('\n')
        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            
            # Check if it's a header
            header_match = re.match(r'^(#+)\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                doc.add_heading(text, level=min(level, 3))
            else:
                # Regular paragraph - remove markdown formatting
                clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                para = doc.add_paragraph(clean_line)
                # Set font size on runs (proper way in python-docx)
                for run in para.runs:
                    run.font.size = Pt(11)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Sanitize filename for safe download
        safe_filename = sanitize_filename(title, story_id, max_length=50)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{safe_filename}_{story_id}.docx"
        )
    except MissingDependencyError:
        raise  # Re-raise MissingDependencyError to be handled by error handler
    except Exception as e:
        logger.error(f"DOCX export failed for story {story_id}: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"DOCX export failed: {str(e)}")


def export_epub(story_text, title, story_id):
    """
    Export story as EPUB (e-book) format.
    
    Requires ebooklib library. Raises MissingDependencyError if not installed.
    
    Args:
        story_text: The story text to export (markdown will be converted to HTML)
        title: The story title (will be sanitized for filename)
        story_id: Unique identifier for the story
    
    Returns:
        Flask send_file response with EPUB attachment
    
    Raises:
        MissingDependencyError: If ebooklib is not installed
        ServiceUnavailableError: If export fails for other reasons
    """
    try:
        from ebooklib import epub  # type: ignore[import-untyped]
    except ImportError:
        raise MissingDependencyError("ebooklib", "pip install ebooklib")
    
    try:
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(f"story_{story_id}")
        book.set_title(title)
        book.set_language('en')
        book.add_author('Short Story Pipeline')
        
        # Remove markdown formatting for EPUB
        text = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        text = text.replace('\n', '<br/>')
        
        # Create chapter
        chapter = epub.EpubHtml(title=title, file_name='chapter.xhtml', lang='en')
        chapter.content = f'<h1>{title}</h1><p>{text}</p>'
        
        # Add chapter
        book.add_item(chapter)
        
        # Add default NCX and Nav file
        book.toc = [chapter]
        book.spine = ['nav', chapter]
        
        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Save to buffer
        buffer = BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        
        # Sanitize filename for safe download
        safe_filename = sanitize_filename(title, story_id, max_length=50)
        
        return send_file(
            buffer,
            mimetype='application/epub+zip',
            as_attachment=True,
            download_name=f"{safe_filename}_{story_id}.epub"
        )
    except MissingDependencyError:
        raise  # Re-raise MissingDependencyError to be handled by error handler
    except Exception as e:
        logger.error(f"EPUB export failed for story {story_id}: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"EPUB export failed: {str(e)}")


if __name__ == '__main__':
    # Production settings
    debug_mode = os.getenv('FLASK_ENV') == 'development'
    port = int(os.getenv('PORT', 5000))
    host = os.getenv('HOST', '0.0.0.0')
    
    app.run(debug=debug_mode, host=host, port=port)

