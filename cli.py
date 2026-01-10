#!/usr/bin/env python3
"""
CLI tool for local story management.

Provides commands for listing, deleting, exporting, and validating stories
without needing to use the web UI.
"""

import os
import sys
from pathlib import Path
from typing import Optional, Dict, Any
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent / "src"))

try:
    import click
except ImportError:
    print("Error: click is required. Install it with: pip install click")
    sys.exit(1)

from src.shortstory.utils import (
    create_story_repository,
    check_distinctiveness,
    validate_premise,
    validate_story_voices,
    MAX_WORD_COUNT,
)
from src.shortstory.exports import (
    export_pdf,
    export_markdown,
    export_txt,
    export_docx,
    export_epub,
    sanitize_filename,
)
from src.shortstory.pipeline import ShortStoryPipeline
from src.shortstory.utils.errors import NotFoundError


# Initialize repository
story_repository = create_story_repository()


def get_story_body(story: Dict[str, Any]) -> str:
    """
    Extract story body text from story dict.
    
    Args:
        story: Story dictionary containing story data. May have 'body' field
               or legacy 'text' field with embedded story content.
    
    Returns:
        The story body text as a string. Returns empty string if no body found.
    """
    if "body" in story:
        return story.get("body", "")
    # Legacy format: extract from text field
    if "text" in story:
        text = story.get("text", "")
        import re
        story_match = re.search(r'## Story\s*\n\s*\n(.+)$', text, re.DOTALL)
        if story_match:
            return story_match.group(1).strip()
        return text
    return ""


def get_story_text(story: Dict[str, Any]) -> str:
    """
    Get full markdown composite text from story.
    
    Generates a composite markdown document containing the story title, genre,
    metadata (tone, pace, POV), character description, theme, and story body.
    
    Args:
        story: Story dictionary containing premise, genre, metadata, and body.
    
    Returns:
        A formatted markdown string containing all story information.
    """
    # If text field exists and body doesn't, use legacy text
    if "text" in story and "body" not in story:
        return story.get("text", "")
    # Generate composite from body + metadata (simplified version)
    body = get_story_body(story)
    metadata = story.get("metadata", {})
    premise = story.get("premise", {})
    genre = story.get("genre", "General Fiction")
    
    idea = premise.get("idea", "") if isinstance(premise, dict) else str(premise) if premise else ""
    character = premise.get("character", {}) if isinstance(premise, dict) else {}
    theme = premise.get("theme", "") if isinstance(premise, dict) else ""
    char_desc = character.get("description", "") if isinstance(character, dict) else str(character) if character else ""
    
    tone = metadata.get('tone', 'balanced')
    pace = metadata.get('pace', 'moderate')
    pov = metadata.get('pov', 'flexible')
    
    return f"""# {idea}

## Genre: {genre}
**Tone:** {tone} | **Pace:** {pace} | **POV:** {pov}

## Character
{char_desc if char_desc else 'Not specified'}

## Theme
{theme if theme else 'Not specified'}

## Story

{body}
"""


@click.group()
def cli():
    """CLI tool for local story management."""
    pass


@cli.command()
@click.option('--page', default=1, type=int, help='Page number (default: 1)')
@click.option('--per-page', default=50, type=int, help='Items per page (default: 50)')
@click.option('--genre', type=str, help='Filter by genre')
@click.option('--format', 'output_format', type=click.Choice(['table', 'json', 'simple']), 
              default='table', help='Output format (default: table)')
def list_stories(page: int, per_page: int, genre: Optional[str], output_format: str) -> None:
    """
    List all stories with metadata.
    
    Displays stories in the specified format (table, json, or simple).
    Supports pagination and optional genre filtering.
    
    Args:
        page: Page number to display (1-indexed).
        per_page: Number of stories per page.
        genre: Optional genre filter. If provided, only stories of this genre are shown.
        output_format: Output format - 'table' for formatted table, 'json' for JSON,
                      or 'simple' for one-line per story.
    
    Raises:
        SystemExit: Exits with code 1 if an error occurs while listing stories.
    """
    try:
        result = story_repository.list(page=page, per_page=per_page, genre=genre)
        stories = result.get('stories', [])
        pagination = result.get('pagination', {})
        
        if output_format == 'json':
            import json
            click.echo(json.dumps(result, indent=2))
        elif output_format == 'simple':
            for story in stories:
                story_id = story.get('id', 'unknown')
                premise = story.get('premise', {})
                idea = premise.get('idea', '') if isinstance(premise, dict) else str(premise) if premise else ''
                genre_name = story.get('genre', 'Unknown')
                word_count = story.get('word_count', 0)
                click.echo(f"{story_id}: {idea[:50]}... ({genre_name}, {word_count} words)")
        else:  # table format
            if not stories:
                click.echo("No stories found.")
                return
            
            # Print header
            click.echo(f"\n{'ID':<12} {'Genre':<20} {'Word Count':<12} {'Premise':<50}")
            click.echo("-" * 100)
            
            for story in stories:
                story_id = story.get('id', 'unknown')
                premise = story.get('premise', {})
                idea = premise.get('idea', '') if isinstance(premise, dict) else str(premise) if premise else ''
                if len(idea) > 48:
                    idea = idea[:45] + "..."
                genre_name = story.get('genre', 'Unknown')
                word_count = story.get('word_count', 0)
                click.echo(f"{story_id:<12} {genre_name:<20} {word_count:<12} {idea:<50}")
            
            # Print pagination info
            total = pagination.get('total', len(stories))
            total_pages = pagination.get('total_pages', 1)
            click.echo(f"\nTotal: {total} stories (Page {page}/{total_pages})")
            
    except Exception as e:
        click.echo(f"Error listing stories: {e}", err=True)
        sys.exit(1)


@cli.command()
@click.argument('story_id')
@click.option('--confirm/--no-confirm', default=False, help='Skip confirmation prompt')
def delete_story(story_id: str, confirm: bool) -> None:
    """
    Delete a story by ID.
    
    Loads the story first to display its information, then prompts for confirmation
    (unless --confirm flag is used) before deleting.
    
    Args:
        story_id: The unique identifier of the story to delete.
        confirm: If True, skip the confirmation prompt and delete immediately.
    
    Raises:
        SystemExit: Exits with code 1 if story not found or deletion fails.
    """
    # Load story first to show what will be deleted
    story = story_repository.load(story_id)
    if not story:
        click.echo(f"Error: Story '{story_id}' not found.", err=True)
        sys.exit(1)
    
    # Show story info
    premise = story.get('premise', {})
    idea = premise.get('idea', '') if isinstance(premise, dict) else str(premise) if premise else ''
    genre = story.get('genre', 'Unknown')
    word_count = story.get('word_count', 0)
    
    click.echo(f"Story ID: {story_id}")
    click.echo(f"Premise: {idea}")
    click.echo(f"Genre: {genre}")
    click.echo(f"Word Count: {word_count}")
    
    # Confirm deletion
    if not confirm:
        if not click.confirm('\nAre you sure you want to delete this story?'):
            click.echo("Deletion cancelled.")
            return
    
    # Delete story
    try:
        success = story_repository.delete(story_id)
        if success:
            click.echo(f"✓ Successfully deleted story '{story_id}'")
        else:
            click.echo(f"Error: Failed to delete story '{story_id}'", err=True)
            sys.exit(1)
    except Exception as e:
        click.echo(f"Error deleting story: {e}", err=True)
        sys.exit(1)


@cli.command()
@click.argument('story_id')
@click.argument('format_type', type=click.Choice(['pdf', 'markdown', 'txt', 'docx', 'epub']))
@click.option('--output', '-o', type=click.Path(), help='Output file path (default: auto-generated)')
def export_story(story_id: str, format_type: str, output: Optional[str]) -> None:
    """
    Export a story to a file in the specified format.
    
    Supports multiple export formats: PDF, Markdown, plain text, DOCX, and EPUB.
    If no output path is specified, a filename is auto-generated from the story title.
    
    Args:
        story_id: The unique identifier of the story to export.
        format_type: Export format - one of 'pdf', 'markdown', 'txt', 'docx', or 'epub'.
        output: Optional output file path. If not provided, a filename is generated
                using the story title and ID.
    
    Raises:
        SystemExit: Exits with code 1 if story not found, has no content, or export fails.
    """
    # Load story
    story = story_repository.load(story_id)
    if not story:
        click.echo(f"Error: Story '{story_id}' not found.", err=True)
        sys.exit(1)
    
    # Get story text
    story_text = get_story_text(story)
    if not story_text:
        click.echo(f"Error: Story '{story_id}' has no content to export.", err=True)
        sys.exit(1)
    
    # Extract title
    import re
    title_match = re.search(r'^#\s+(.+)$', story_text, re.MULTILINE)
    if title_match:
        raw_title = title_match.group(1)
    else:
        premise = story.get("premise", {})
        if isinstance(premise, dict):
            raw_title = premise.get("idea", f"Story {story_id}")
        else:
            raw_title = f"Story {story_id}"
    
    # Clean title for document content (remove dangerous chars but keep longer for readability)
    title = re.sub(r'[<>"\';\\/]', '', raw_title).strip()
    if not title:
        title = f"Story {story_id}"
    
    # Generate output filename if not provided using centralized sanitization
    if not output:
        safe_title = sanitize_filename(title, story_id, max_length=50)
        ext_map = {
            'pdf': '.pdf',
            'markdown': '.md',
            'txt': '.txt',
            'docx': '.docx',
            'epub': '.epub'
        }
        output = f"{safe_title}_{story_id}{ext_map[format_type]}"
    
    try:
        # Export to file
        if format_type == 'pdf':
            # PDF export returns Flask response, we need to handle it differently
            from io import BytesIO
            buffer = BytesIO()
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.colors import HexColor
            
            doc = SimpleDocTemplate(buffer, pagesize=letter,
                                   rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=18)
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=HexColor('#667eea'),
                spaceAfter=30,
                alignment=TA_LEFT
            )
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                leading=16,
                spaceAfter=12,
                alignment=TA_LEFT
            )
            
            story_content = []
            story_content.append(Paragraph(title, title_style))
            story_content.append(Spacer(1, 0.2*inch))
            
            lines = story_text.split('\n')
            for line in lines:
                if line.strip():
                    clean_line = re.sub(r'^#+\s+', '', line)
                    clean_line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean_line)
                    clean_line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', clean_line)
                    story_content.append(Paragraph(clean_line, body_style))
                else:
                    story_content.append(Spacer(1, 0.1*inch))
            
            doc.build(story_content)
            buffer.seek(0)
            
            with open(output, 'wb') as f:
                f.write(buffer.read())
                
        elif format_type == 'markdown':
            with open(output, 'w', encoding='utf-8') as f:
                f.write(story_text)
                
        elif format_type == 'txt':
            # Remove markdown formatting
            text = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            with open(output, 'w', encoding='utf-8') as f:
                f.write(text)
                
        elif format_type == 'docx':
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                title_para = doc.add_heading(title, level=1)
                title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                lines = story_text.split('\n')
                for line in lines:
                    if not line.strip():
                        doc.add_paragraph()
                        continue
                    
                    header_match = re.match(r'^(#+)\s+(.+)$', line)
                    if header_match:
                        level = len(header_match.group(1))
                        text = header_match.group(2)
                        doc.add_heading(text, level=min(level, 3))
                    else:
                        clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                        clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                        para = doc.add_paragraph(clean_line)
                        for run in para.runs:
                            run.font.size = Pt(11)
                
                doc.save(output)
            except ImportError:
                click.echo("Error: python-docx is required for DOCX export. Install it with: pip install python-docx", err=True)
                sys.exit(1)
                
        elif format_type == 'epub':
            try:
                from ebooklib import epub
                
                book = epub.EpubBook()
                book.set_identifier(f"story_{story_id}")
                book.set_title(title)
                book.set_language('en')
                book.add_author('Short Story Pipeline')
                
                text = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
                text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
                text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
                text = text.replace('\n', '<br/>')
                
                chapter = epub.EpubHtml(title=title, file_name='chapter.xhtml', lang='en')
                chapter.content = f'<h1>{title}</h1><p>{text}</p>'
                
                book.add_item(chapter)
                book.toc = [chapter]
                book.spine = ['nav', chapter]
                book.add_item(epub.EpubNcx())
                book.add_item(epub.EpubNav())
                
                epub.write_epub(output, book, {})
            except ImportError:
                click.echo("Error: ebooklib is required for EPUB export. Install it with: pip install ebooklib", err=True)
                sys.exit(1)
        
        click.echo(f"✓ Successfully exported story '{story_id}' to '{output}' ({format_type.upper()})")
        
    except Exception as e:
        click.echo(f"Error exporting story: {e}", err=True)
        import traceback
        if os.getenv('DEBUG'):
            traceback.print_exc()
        sys.exit(1)


@cli.command()
@click.argument('story_id')
@click.option('--verbose', '-v', is_flag=True, help='Show detailed validation results')
def validate_story(story_id: str, verbose: bool) -> None:
    """
    Validate a story's content, word count, and distinctiveness.
    
    Performs comprehensive validation including:
    - Word count validation against the maximum limit
    - Distinctiveness scoring for idea and character
    - Voice consistency (if verbose mode enabled)
    - Premise validation (if verbose mode enabled)
    
    Args:
        story_id: The unique identifier of the story to validate.
        verbose: If True, includes additional validation checks for voice consistency
                and premise validation.
    
    Raises:
        SystemExit: Exits with code 1 if story not found or has no content.
    """
    # Load story
    story = story_repository.load(story_id)
    if not story:
        click.echo(f"Error: Story '{story_id}' not found.", err=True)
        sys.exit(1)
    
    # Get story body
    body = get_story_body(story)
    if not body:
        click.echo(f"Error: Story '{story_id}' has no content to validate.", err=True)
        sys.exit(1)
    
    click.echo(f"Validating story: {story_id}\n")
    
    # Word count validation
    try:
        pipeline = ShortStoryPipeline()
        word_count, is_valid = pipeline.word_validator.validate(body, raise_error=False)
        
        click.echo("Word Count Validation:")
        click.echo(f"  Words: {word_count:,} / {MAX_WORD_COUNT:,}")
        click.echo(f"  Remaining: {MAX_WORD_COUNT - word_count:,}")
        if is_valid:
            click.echo(f"  Status: ✓ Valid")
        else:
            click.echo(f"  Status: ✗ Exceeds limit")
    except Exception as e:
        click.echo(f"  Status: ✗ Error: {e}")
    
    # Distinctiveness validation
    try:
        premise = story.get('premise', {})
        idea = premise.get('idea', '') if isinstance(premise, dict) else str(premise) if premise else ''
        character = premise.get('character', {}) if isinstance(premise, dict) else {}
        
        click.echo("\nDistinctiveness Validation:")
        
        if idea:
            idea_dist = check_distinctiveness(idea)
            idea_score = idea_dist.get('distinctiveness_score', 0.0) if isinstance(idea_dist, dict) else 0.0
            has_cliches = idea_dist.get('has_cliches', False) if isinstance(idea_dist, dict) else False
            cliches = idea_dist.get('found_cliches', []) if isinstance(idea_dist, dict) else []
            
            click.echo(f"  Idea Score: {idea_score:.2f}/1.0")
            if has_cliches:
                click.echo(f"  Clichés Found: {', '.join(cliches) if cliches else 'Yes (unspecified)'}")
            else:
                click.echo(f"  Clichés: ✓ None detected")
        
        if character:
            char_dist = check_distinctiveness(None, character=character)
            char_score = char_dist.get('distinctiveness_score', 0.0) if isinstance(char_dist, dict) else 0.0
            has_generic = char_dist.get('has_generic_archetype', False) if isinstance(char_dist, dict) else False
            generic = char_dist.get('generic_elements', []) if isinstance(char_dist, dict) else []
            
            click.echo(f"  Character Score: {char_score:.2f}/1.0")
            if has_generic:
                click.echo(f"  Generic Elements: {', '.join(generic) if generic else 'Yes (unspecified)'}")
            else:
                click.echo(f"  Generic Elements: ✓ None detected")
    except Exception as e:
        click.echo(f"  Error: {e}")
    
    # Voice validation (if verbose)
    if verbose:
        try:
            click.echo("\nVoice Validation:")
            character_info = premise.get('character', {}) if isinstance(premise, dict) else {}
            voice_result = validate_story_voices(body, character_info if character_info else None)
            
            if isinstance(voice_result, dict):
                is_consistent = voice_result.get('is_consistent', False)
                issues = voice_result.get('issues', [])
                
                click.echo(f"  Consistency: {'✓ Consistent' if is_consistent else '✗ Issues found'}")
                if issues:
                    for issue in issues:
                        click.echo(f"    - {issue}")
            else:
                click.echo(f"  Result: {voice_result}")
        except Exception as e:
            click.echo(f"  Error: {e}")
    
    # Premise validation (if verbose)
    if verbose:
        try:
            click.echo("\nPremise Validation:")
            premise_result = validate_premise(idea, character, premise.get('theme', '') if isinstance(premise, dict) else '')
            
            if isinstance(premise_result, dict):
                is_valid_premise = premise_result.get('is_valid', False)
                warnings = premise_result.get('warnings', [])
                
                click.echo(f"  Valid: {'✓ Yes' if is_valid_premise else '✗ No'}")
                if warnings:
                    for warning in warnings:
                        click.echo(f"    - {warning}")
            else:
                click.echo(f"  Result: {premise_result}")
        except Exception as e:
            click.echo(f"  Error: {e}")
    
    click.echo("\n✓ Validation complete")


if __name__ == '__main__':
    cli()

