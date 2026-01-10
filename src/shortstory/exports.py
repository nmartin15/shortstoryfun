"""
Export functions for stories in various formats.

All export functions return Flask responses suitable for download.
"""

import re
import logging
from typing import Dict, Any
from io import BytesIO
from flask import send_file, Response

from .utils.errors import (
    ValidationError,
    MissingDependencyError,
    ServiceUnavailableError
)

logger = logging.getLogger(__name__)

# Pre-compiled regex patterns for filename sanitization (performance optimization)
# These are compiled once at module load time to avoid recompilation on every call
_PATH_TRAVERSAL_PATTERN = re.compile(r'\.\.+')
_DANGEROUS_CHARS_PATTERN = re.compile(r'[<>:"|?*\\/;&`$]')
_SCRIPT_TAG_PATTERN = re.compile(r'<script[^>]*>.*?</script>', re.IGNORECASE | re.DOTALL)
_JAVASCRIPT_PROTOCOL_PATTERN = re.compile(r'javascript:', re.IGNORECASE)
_EVENT_HANDLER_PATTERN = re.compile(r'on\w+\s*=', re.IGNORECASE)
_WHITESPACE_PATTERN = re.compile(r'\s+')
_NON_ALPHANUMERIC_PATTERN = re.compile(r'[^a-zA-Z0-9_-]')


def sanitize_filename(title: str, story_id: str, max_length: int = 50) -> str:
    """
    Sanitize a title for use in filenames.
    
    Removes all dangerous characters that could lead to:
    - Path traversal attacks (.., /, \)
    - Command injection (|, &, ;, `)
    - OS-specific issues (:, *, ?, <, >, ")
    - XSS in download attributes (script tags, HTML entities)
    
    Args:
        title: Original title
        story_id: Story ID for fallback
        max_length: Maximum length for filename
        
    Returns:
        Sanitized filename-safe string containing only alphanumeric, underscores, and hyphens
    """
    if not title:
        return f"Story_{story_id[:8]}"
    
    # Remove path traversal sequences first (before other sanitization)
    # Use pre-compiled pattern for better performance
    safe = _PATH_TRAVERSAL_PATTERN.sub('', title)
    
    # Remove all dangerous characters:
    # - Path separators: /, \
    # - Windows forbidden: <, >, :, ", |, ?, *
    # - Shell metacharacters: &, ;, `, $
    # - HTML/script tags: <, >
    # - Quotes: ", '
    # Use pre-compiled pattern for better performance
    safe = _DANGEROUS_CHARS_PATTERN.sub('', safe)
    
    # Remove script tags and common XSS patterns
    # Use pre-compiled patterns for better performance
    safe = _SCRIPT_TAG_PATTERN.sub('', safe)
    safe = _JAVASCRIPT_PROTOCOL_PATTERN.sub('', safe)
    safe = _EVENT_HANDLER_PATTERN.sub('', safe)  # Remove event handlers like onclick=
    
    # Replace spaces and other whitespace with underscores
    safe = _WHITESPACE_PATTERN.sub('_', safe)
    
    # Remove any remaining non-alphanumeric except underscores and hyphens
    # This is a final safety net to catch any edge cases
    safe = _NON_ALPHANUMERIC_PATTERN.sub('', safe)
    
    # Remove leading/trailing underscores and hyphens (can cause issues on some systems)
    safe = safe.strip('_-')
    
    # Truncate if too long
    if len(safe) > max_length:
        safe = safe[:max_length]
    
    # Fallback if empty or only dangerous characters
    if not safe:
        # Sanitize story_id as well for fallback
        # Use pre-compiled pattern for better performance
        safe_id = _NON_ALPHANUMERIC_PATTERN.sub('', story_id)[:8]
        safe = f"Story_{safe_id}" if safe_id else "Story_export"
    
    return safe


def export_pdf(story_text: str, title: str, story_id: str) -> Response:
    """
    Export story as PDF.
    
    Args:
        story_text: Full story text (markdown format)
        title: Story title
        story_id: Story identifier
        
    Returns:
        Flask response with PDF file
        
    Raises:
        MissingDependencyError: If reportlab is not installed
        ServiceUnavailableError: If PDF generation fails
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.colors import HexColor
    except ImportError:
        raise MissingDependencyError("reportlab", "pip install reportlab")
    
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=HexColor('#667eea'),
            spaceAfter=30,
            alignment=TA_LEFT
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        story_content = []
        story_content.append(Paragraph(title, title_style))
        story_content.append(Spacer(1, 0.2 * inch))
        
        lines = story_text.split('\n')
        for line in lines:
            if line.strip():
                clean_line = re.sub(r'^#+\s+', '', line)
                clean_line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean_line)
                clean_line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', clean_line)
                story_content.append(Paragraph(clean_line, body_style))
            else:
                story_content.append(Spacer(1, 0.1 * inch))
        
        doc.build(story_content)
        buffer.seek(0)
        
        safe_filename = sanitize_filename(title, story_id, max_length=50)
        # Use pre-compiled pattern for better performance
        sanitized_story_id = _NON_ALPHANUMERIC_PATTERN.sub('', story_id)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{safe_filename}_{sanitized_story_id}.pdf"
        )
    except (IOError, OSError) as e:
        logger.error(f"I/O error during PDF export: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"PDF export failed due to I/O issue: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error during PDF export: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"PDF export failed: {str(e)}")


def export_markdown(story_text: str, title: str, story_id: str) -> Response:
    """
    Export story as Markdown.
    
    Args:
        story_text: Full story text (markdown format)
        title: Story title
        story_id: Story identifier
        
    Returns:
        Flask response with Markdown file
    """
    buffer = BytesIO()
    buffer.write(story_text.encode('utf-8'))
    buffer.seek(0)
    
    safe_filename = sanitize_filename(title, story_id, max_length=50)
    sanitized_story_id = re.sub(r'[^a-zA-Z0-9_-]', '', story_id)
    
    return send_file(
        buffer,
        mimetype='text/markdown',
        as_attachment=True,
        download_name=f"{safe_filename}_{sanitized_story_id}.md"
    )


def export_txt(story_text: str, title: str, story_id: str) -> Response:
    """
    Export story as plain text (markdown formatting removed).
    
    Args:
        story_text: Full story text (markdown format)
        title: Story title
        story_id: Story identifier
        
    Returns:
        Flask response with text file
    """
    # Remove markdown formatting
    text = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    
    buffer = BytesIO()
    buffer.write(text.encode('utf-8'))
    buffer.seek(0)
    
    safe_filename = sanitize_filename(title, story_id, max_length=50)
    sanitized_story_id = re.sub(r'[^a-zA-Z0-9_-]', '', story_id)
    
    return send_file(
        buffer,
        mimetype='text/plain',
        as_attachment=True,
        download_name=f"{safe_filename}_{sanitized_story_id}.txt"
    )


def export_docx(story_text: str, title: str, story_id: str) -> Response:
    """
    Export story as DOCX.
    
    Args:
        story_text: Full story text (markdown format)
        title: Story title
        story_id: Story identifier
        
    Returns:
        Flask response with DOCX file
        
    Raises:
        MissingDependencyError: If python-docx is not installed
        ServiceUnavailableError: If DOCX generation fails
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise MissingDependencyError("python-docx", "pip install python-docx")
    
    try:
        doc = Document()
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        lines = story_text.split('\n')
        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                continue
            
            header_match = re.match(r'^(#+)\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                doc.add_heading(text, level=min(level, 3))
            else:
                clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                para = doc.add_paragraph(clean_line)
                for run in para.runs:
                    run.font.size = Pt(11)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        safe_filename = sanitize_filename(title, story_id, max_length=50)
        # Use pre-compiled pattern for better performance
        sanitized_story_id = _NON_ALPHANUMERIC_PATTERN.sub('', story_id)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{safe_filename}_{sanitized_story_id}.docx"
        )
    except (IOError, OSError) as e:
        logger.error(f"I/O error during DOCX export: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"DOCX export failed due to I/O issue: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error during DOCX export: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"DOCX export failed: {str(e)}")


def export_epub(story_text: str, title: str, story_id: str) -> Response:
    """
    Export story as EPUB.
    
    Args:
        story_text: Full story text (markdown format)
        title: Story title
        story_id: Story identifier
        
    Returns:
        Flask response with EPUB file
        
    Raises:
        MissingDependencyError: If ebooklib is not installed
        ServiceUnavailableError: If EPUB generation fails
    """
    try:
        import ebooklib
        from ebooklib import epub
    except ImportError:
        raise MissingDependencyError("ebooklib", "pip install ebooklib")
    
    try:
        book = epub.EpubBook()
        # Sanitize story_id for EPUB identifier
        sanitized_id = re.sub(r'[^a-zA-Z0-9_-]', '', story_id)
        book.set_identifier(f"story_{sanitized_id}")
        book.set_title(title)
        book.set_language('en')
        book.add_author('Short Story Pipeline')
        
        text = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        text = text.replace('\n', '<br/>')
        
        chapter = epub.EpubHtml(title=title, file_name='chapter.xhtml', lang='en')
        chapter.content = f'<h1>{title}</h1><p>{text}</p>'
        
        book.add_item(chapter)
        book.toc = [chapter]
        book.spine = ['nav', chapter]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        buffer = BytesIO()
        epub.write_epub(buffer, book, {})
        buffer.seek(0)
        
        safe_filename = sanitize_filename(title, story_id, max_length=50)
        # Use pre-compiled pattern for better performance
        sanitized_story_id = _NON_ALPHANUMERIC_PATTERN.sub('', story_id)
        
        return send_file(
            buffer,
            mimetype='application/epub+zip',
            as_attachment=True,
            download_name=f"{safe_filename}_{sanitized_story_id}.epub"
        )
    except (IOError, OSError) as e:
        logger.error(f"I/O error during EPUB export: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"EPUB export failed due to I/O issue: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error during EPUB export: {str(e)}", exc_info=True)
        raise ServiceUnavailableError("export", f"EPUB export failed: {str(e)}")


def export_story_from_dict(
    story: Dict[str, Any],
    story_id: str,
    format_type: str,
    story_text: str
) -> Response:
    """
    Export a story in the specified format.
    
    This is the main entry point for story exports from Flask routes.
    
    Args:
        story: Story dictionary
        story_id: Story identifier
        format_type: Export format (pdf, markdown, txt, docx, epub)
        story_text: Full story text (composite markdown)
        
    Returns:
        Flask response with exported file
        
    Raises:
        ValidationError: If format is invalid or story has no content
        MissingDependencyError: If required library is not installed
        ServiceUnavailableError: If export fails
    """
    # Validate format
    valid_formats = ['pdf', 'markdown', 'txt', 'docx', 'epub']
    if format_type not in valid_formats:
        raise ValidationError(
            f"Invalid format '{format_type}'. Supported formats: {', '.join(valid_formats)}",
            details={"format_type": format_type, "valid_formats": valid_formats}
        )
    
    # Validate story has content
    if not story_text or not story_text.strip():
        raise ValidationError(
            "Story has no content to export.",
            details={"story_id": story_id}
        )
    
    # Extract title from story text or metadata
    title_match = re.search(r'^#\s+(.+)$', story_text, re.MULTILINE)
    if title_match:
        raw_title = title_match.group(1)
    else:
        premise = story.get("premise", {})
        if isinstance(premise, dict):
            raw_title = premise.get("idea", f"Story {story_id}")
        else:
            raw_title = f"Story {story_id}"
    
    # Clean title
    title = re.sub(r'[<>"\';\\/]', '', raw_title).strip()
    if not title:
        title = f"Story {story_id}"
    
    # Route to appropriate export function
    try:
        if format_type == 'pdf':
            return export_pdf(story_text, title, story_id)
        elif format_type == 'markdown':
            return export_markdown(story_text, title, story_id)
        elif format_type == 'txt':
            return export_txt(story_text, title, story_id)
        elif format_type == 'docx':
            return export_docx(story_text, title, story_id)
        elif format_type == 'epub':
            return export_epub(story_text, title, story_id)
    except (ValidationError, MissingDependencyError, ServiceUnavailableError):
        # Re-raise known error types to be handled by error handlers
        raise
    except (IOError, OSError) as e:
        logger.error(
            f"I/O error during export for story {story_id}, format {format_type}: {str(e)}",
            exc_info=True
        )
        raise ServiceUnavailableError("export", f"Export failed due to I/O issue: {str(e)}")
    except Exception as e:
        logger.error(
            f"Unexpected export failure for story {story_id}, format {format_type}: {str(e)}",
            exc_info=True
        )
        raise ServiceUnavailableError(
            "export",
            f"An unexpected error occurred during export: {str(e)}"
        )
