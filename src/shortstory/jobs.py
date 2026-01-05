"""
Background job tasks for ShortStory application.

This module contains all background tasks that can be executed asynchronously
using RQ (Redis Queue). Heavy operations like LLM calls, story generation,
revisions, and batch exports are moved here to prevent blocking the web server.
"""

import logging
import traceback
from typing import Dict, Any, Optional
from datetime import datetime

from .pipeline import ShortStoryPipeline
from .genres import get_genre_config
from .utils import check_distinctiveness
from .exports import export_pdf, export_markdown, export_txt, export_docx, export_epub
from .utils.storage import get_storage
from .utils.errors import ValidationError, MissingDependencyError

logger = logging.getLogger(__name__)


def generate_story_job(
    idea: str,
    character: Dict[str, Any],
    theme: str,
    genre: str = "General Fiction",
    max_word_count: int = 7500
) -> Dict[str, Any]:
    """
    Background job for generating a complete story.
    
    This is a long-running task that performs all pipeline stages:
    - Premise capture
    - Outline generation
    - Scaffolding
    - Drafting
    - Revision
    
    Args:
        idea: Story idea/premise
        character: Character description (dict)
        theme: Central theme
        genre: Genre name
        max_word_count: Maximum word count
        
    Returns:
        Dict containing:
            - status: "completed" or "failed"
            - story_id: Generated story ID
            - story_data: Complete story data (if successful)
            - error: Error message (if failed)
    """
    try:
        logger.info(f"Starting story generation job: idea='{idea[:50]}...', genre={genre}")
        
        if not idea or not idea.strip():
            raise ValidationError("Story idea is required")
        
        if isinstance(character, str):
            character = {"description": character}
        
        genre_config = get_genre_config(genre)
        if genre_config is None:
            genre_config = get_genre_config('General Fiction')
        if genre_config is None:
            raise ValueError("Genre configuration service unavailable")
        
        # Create pipeline instance
        pipeline = ShortStoryPipeline(
            max_word_count=max_word_count,
            genre=genre,
            genre_config=genre_config
        )
        
        # Run full pipeline
        premise = pipeline.capture_premise(idea, character, theme, validate=True)
        outline = pipeline.generate_outline(genre=genre)
        scaffold = pipeline.scaffold(genre=genre)
        draft = pipeline.draft()
        revised_draft = pipeline.revise()
        
        # Extract metadata
        constraints = genre_config.get('constraints', {})
        tone = scaffold.get('tone', constraints.get('tone', 'balanced')) if isinstance(scaffold, dict) else constraints.get('tone', 'balanced')
        pace = scaffold.get('pace', constraints.get('pace', 'moderate')) if isinstance(scaffold, dict) else constraints.get('pace', 'moderate')
        pov = scaffold.get('pov', constraints.get('pov_preference', 'flexible')) if isinstance(scaffold, dict) else constraints.get('pov_preference', 'flexible')
        
        idea_dist = check_distinctiveness(idea)
        char_dist = check_distinctiveness(None, character=character)
        
        # Get revised story text
        revised_story_text = revised_draft.get('text', '')
        story_word_count = pipeline.word_validator.count_words(revised_story_text)
        
        # Generate story ID
        import uuid
        story_id = f"story_{uuid.uuid4().hex[:8]}"
        
        # Build story metadata
        story_metadata = {
            "tone": tone,
            "pace": pace,
            "pov": pov,
            "idea_distinctiveness": idea_dist,
            "character_distinctiveness": char_dist
        }
        
        # Build standardized story data using story_builder
        # This ensures consistency across all story creation points
        from src.shortstory.utils.story_builder import build_story_data
        
        story_data = build_story_data(
            story_id=story_id,
            premise=premise,
            outline=outline,
            genre=genre,
            genre_config=genre_config,
            body=revised_story_text,
            word_count=story_word_count,
            scaffold=scaffold,
            metadata=story_metadata,
            draft=draft,
            revised_draft=revised_draft,
            max_words=max_word_count
        )
        
        # Store story
        storage = get_storage()
        storage.save_story(story_id, story_data)
        
        logger.info(f"Story generation completed: story_id={story_id}, word_count={story_word_count}")
        
        return {
            "status": "completed",
            "story_id": story_id,
            "story_data": story_data,
            "word_count": story_word_count
        }
        
    except Exception as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()
        logger.error(f"Story generation job failed: {error_msg}\n{error_trace}")
        
        return {
            "status": "failed",
            "error": error_msg,
            "error_type": type(e).__name__
        }


def revise_story_job(
    story_id: str,
    use_llm: bool = True
) -> Dict[str, Any]:
    """
    Background job for revising an existing story.
    
    Args:
        story_id: ID of the story to revise
        use_llm: Whether to use LLM for revision
        
    Returns:
        Dict containing:
            - status: "completed" or "failed"
            - story_id: Story ID
            - revised_body: Revised story text (if successful)
            - word_count: New word count (if successful)
            - error: Error message (if failed)
    """
    try:
        logger.info(f"Starting story revision job: story_id={story_id}")
        
        storage = get_storage()
        story = storage.get_story(story_id)
        
        if story is None:
            raise ValueError(f"Story not found: {story_id}")
        
        # Get current story body
        current_body = story.get("body", "")
        if not current_body:
            raise ValidationError("Story has no content to revise")
        
        # Get genre configuration
        story_genre = story.get("genre", "General Fiction")
        story_genre_config = story.get("genre_config")
        if story_genre_config is None:
            story_genre_config = get_genre_config(story_genre)
        
        # Create pipeline instance
        pipeline = ShortStoryPipeline(genre=story_genre, genre_config=story_genre_config)
        
        # Create temporary draft object
        temp_draft = {
            "text": current_body,
            "word_count": story.get("word_count", 0)
        }
        
        # Run revision
        revised_draft = pipeline.revise(draft=temp_draft, use_llm=use_llm)
        revised_body = revised_draft.get('text', '')
        revised_word_count = revised_draft.get('word_count', 0)
        
        # Update story
        story["body"] = revised_body
        story["word_count"] = revised_word_count
        story["max_words"] = story.get("max_words", 7500)
        story["updated_at"] = datetime.now().isoformat()
        
        # Add to revision history
        if "revision_history" not in story:
            story["revision_history"] = []
        if "current_revision" not in story:
            story["current_revision"] = 0
        
        new_version = story["current_revision"] + 1
        story["revision_history"].append({
            "version": new_version,
            "body": revised_body,
            "word_count": revised_word_count,
            "type": "revised",
            "timestamp": datetime.now().isoformat()
        })
        story["current_revision"] = new_version
        
        # Save updated story
        storage.save_story(story_id, story)
        
        logger.info(f"Story revision completed: story_id={story_id}, version={new_version}, word_count={revised_word_count}")
        
        return {
            "status": "completed",
            "story_id": story_id,
            "revised_body": revised_body,
            "word_count": revised_word_count,
            "version": new_version
        }
        
    except Exception as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()
        logger.error(f"Story revision job failed: story_id={story_id}, error={error_msg}\n{error_trace}")
        
        return {
            "status": "failed",
            "story_id": story_id,
            "error": error_msg,
            "error_type": type(e).__name__
        }


def export_story_job(
    story_id: str,
    format_type: str
) -> Dict[str, Any]:
    """
    Background job for exporting a story in various formats.
    
    This is useful for batch exports or large stories that take time to process.
    Note: For most exports, synchronous processing is faster. This is mainly
    useful for batch operations.
    
    Args:
        story_id: ID of the story to export
        format_type: Export format (pdf, markdown, txt, docx, epub)
        
    Returns:
        Dict containing:
            - status: "completed" or "failed"
            - story_id: Story ID
            - format_type: Export format
            - file_size: Size of exported file in bytes (if successful)
            - error: Error message (if failed)
    """
    try:
        logger.info(f"Starting story export job: story_id={story_id}, format={format_type}")
        
        storage = get_storage()
        story = storage.get_story(story_id)
        
        if story is None:
            raise ValueError(f"Story not found: {story_id}")
        
        # Get story text (composite with metadata)
        from .utils.storage import get_story_text
        story_text = get_story_text(story)
        
        if not story_text:
            raise ValidationError("Story has no content to export")
        
        # Extract title
        import re
        title_match = re.search(r'^#\s+(.+)$', story_text, re.MULTILINE)
        if title_match:
            raw_title = title_match.group(1)
        else:
            premise = story.get("premise", {})
            if isinstance(premise, dict):
                raw_title = premise.get("idea", f"Story {story_id}")
            else:
                raw_title = f"Story {story_id}"
        
        # Clean title
        title = re.sub(r'[<>"\';\\/]', '', raw_title).strip()
        if not title:
            title = f"Story {story_id}"
        
        # Generate export content based on format
        # Since export functions return Flask responses, we generate content directly
        import tempfile
        import os
        from io import BytesIO
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{format_type}") as tmp_file:
            tmp_path = tmp_file.name
            
            if format_type == 'pdf':
                # Generate PDF content
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.enums import TA_LEFT
                from reportlab.lib.colors import HexColor
                
                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=letter,
                                       rightMargin=72, leftMargin=72,
                                       topMargin=72, bottomMargin=18)
                
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    textColor=HexColor('#667eea'),
                    spaceAfter=30,
                    alignment=TA_LEFT
                )
                body_style = ParagraphStyle(
                    'CustomBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=16,
                    spaceAfter=12,
                    alignment=TA_LEFT
                )
                
                story_content = []
                story_content.append(Paragraph(title, title_style))
                story_content.append(Spacer(1, 0.2*inch))
                
                # Process story text
                lines = story_text.split('\n')
                for line in lines:
                    if line.strip():
                        # Remove markdown headers
                        clean_line = re.sub(r'^#+\s+', '', line)
                        story_content.append(Paragraph(clean_line, body_style))
                
                doc.build(story_content)
                tmp_file.write(buffer.getvalue())
                
            elif format_type == 'markdown':
                tmp_file.write(story_text.encode('utf-8'))
            elif format_type == 'txt':
                # Strip markdown formatting for plain text
                text_content = re.sub(r'^#+\s+', '', story_text, flags=re.MULTILINE)
                tmp_file.write(text_content.encode('utf-8'))
            elif format_type == 'docx':
                # For DOCX, use python-docx directly
                try:
                    from docx import Document
                    from docx.shared import Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = Document()
                    title_para = doc.add_heading(title, level=1)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Process story text
                    lines = story_text.split('\n')
                    for line in lines:
                        if not line.strip():
                            doc.add_paragraph()
                            continue
                        
                        # Check if it's a header
                        header_match = re.match(r'^(#+)\s+(.+)$', line)
                        if header_match:
                            level = len(header_match.group(1))
                            text = header_match.group(2)
                            doc.add_heading(text, level=min(level, 3))
                        else:
                            # Regular paragraph - remove markdown formatting
                            clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
                            clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
                            para = doc.add_paragraph(clean_line)
                            # Set font size
                            for run in para.runs:
                                run.font.size = Pt(11)
                    
                    doc.save(tmp_file)
                except ImportError:
                    raise MissingDependencyError("python-docx", "pip install python-docx")
            elif format_type == 'epub':
                # EPUB export is complex - for now, just save as text
                # Full implementation would use ebooklib
                logger.warning("EPUB export in background jobs not fully implemented")
                tmp_file.write(story_text.encode('utf-8'))
            else:
                raise ValidationError(f"Unsupported export format: {format_type}")
        
        file_size = os.path.getsize(tmp_path)
        
        logger.info(f"Story export completed: story_id={story_id}, format={format_type}, file={tmp_path}, size={file_size}")
        
        return {
            "status": "completed",
            "story_id": story_id,
            "format_type": format_type,
            "file_path": tmp_path,
            "file_size": file_size,
            "title": title
        }
        
    except Exception as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()
        logger.error(f"Story export job failed: story_id={story_id}, format={format_type}, error={error_msg}\n{error_trace}")
        
        return {
            "status": "failed",
            "story_id": story_id,
            "format_type": format_type,
            "error": error_msg,
            "error_type": type(e).__name__
        }


def batch_export_job(
    story_ids: list,
    format_type: str
) -> Dict[str, Any]:
    """
    Background job for batch exporting multiple stories.
    
    Args:
        story_ids: List of story IDs to export
        format_type: Export format (pdf, markdown, txt, docx, epub)
        
    Returns:
        Dict containing:
            - status: "completed" or "failed"
            - results: List of export results for each story
            - total: Total number of stories
            - succeeded: Number of successful exports
            - failed: Number of failed exports
    """
    try:
        logger.info(f"Starting batch export job: {len(story_ids)} stories, format={format_type}")
        
        results = []
        succeeded = 0
        failed = 0
        
        for story_id in story_ids:
            try:
                result = export_story_job(story_id, format_type)
                results.append({
                    "story_id": story_id,
                    "status": result.get("status"),
                    "file_path": result.get("file_path") if result.get("status") == "completed" else None,
                    "error": result.get("error") if result.get("status") == "failed" else None
                })
                if result.get("status") == "completed":
                    succeeded += 1
                else:
                    failed += 1
            except Exception as e:
                logger.error(f"Failed to export story {story_id}: {str(e)}")
                results.append({
                    "story_id": story_id,
                    "status": "failed",
                    "error": str(e)
                })
                failed += 1
        
        logger.info(f"Batch export completed: {succeeded} succeeded, {failed} failed")
        
        return {
            "status": "completed",
            "results": results,
            "total": len(story_ids),
            "succeeded": succeeded,
            "failed": failed
        }
        
    except Exception as e:
        error_msg = str(e)
        error_trace = traceback.format_exc()
        logger.error(f"Batch export job failed: error={error_msg}\n{error_trace}")
        
        return {
            "status": "failed",
            "error": error_msg,
            "error_type": type(e).__name__
        }

